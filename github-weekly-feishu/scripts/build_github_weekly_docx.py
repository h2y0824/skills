import argparse
import json
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from zoneinfo import ZoneInfo
except ImportError:  # pragma: no cover
    ZoneInfo = None


DEFAULT_OUTPUT_DIR = Path(r"F:\github每周热点\GitHub热门内容推荐")


def beijing_today():
    if ZoneInfo:
        return datetime.now(ZoneInfo("Asia/Shanghai")).date()
    return datetime.now().date()


def format_date(date_value):
    return f"{date_value.year}.{date_value.month}.{date_value.day}"


def previous_week_range(today=None):
    today = today or beijing_today()
    current_monday = today - timedelta(days=today.weekday())
    start = current_monday - timedelta(days=7)
    end = current_monday - timedelta(days=1)
    return start, end


def default_output_path(today=None):
    start, end = previous_week_range(today)
    filename = f"{format_date(start)}-{end.month}.{end.day}GitHub热门内容推荐.docx"
    return DEFAULT_OUTPUT_DIR / filename


def set_run(run, size=10.5, bold=False, color=None):
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text="", size=10.5, bold=False, color=None, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.name = "Microsoft YaHei"
    run.font.color.rgb = RGBColor(20, 39, 54)
    run.font.size = Pt(18 if level == 1 else 14)
    run.bold = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)


def add_image(doc, image_path, caption):
    if not image_path:
        return
    path = Path(image_path)
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(path), width=Inches(6.25))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(10)
    run = c.add_run(caption or path.name)
    set_run(run, size=9, color=(95, 104, 113))


def add_bullets(doc, items):
    for item in items or []:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run(run, size=10)


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["项目", "方向", "主要价值", "推荐理由"]
    for i, header in enumerate(headers):
        run = table.rows[0].cells[i].paragraphs[0].add_run(header)
        set_run(run, size=9.5, bold=True, color=(20, 39, 54))
    for row in rows:
        cells = table.add_row().cells
        values = [row.get("name", ""), row.get("direction", ""), row.get("value", ""), row.get("reason", "")]
        for i, value in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run(run, size=8.5)


def build_docx(article, out_path):
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run(article["title"])
    set_run(run, size=24, bold=True, color=(20, 39, 54))

    subtitle_text = article.get("subtitle")
    if subtitle_text:
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(14)
        run = subtitle.add_run(subtitle_text)
        set_run(run, size=13, color=(12, 111, 126))

    cover = article.get("cover_image")
    if cover:
        add_image(doc, cover, article.get("cover_caption", "封面图"))

    if article.get("metadata"):
        add_para(doc, article["metadata"], size=9, color=(95, 104, 113))

    for paragraph in article.get("intro", []):
        add_para(doc, paragraph)

    add_heading(doc, "本期项目总览", 1)
    add_table(doc, article.get("summary", []))

    for index, project in enumerate(article.get("projects", []), start=1):
        add_heading(doc, f"{index}. {project['name']}", 1)
        add_image(doc, project.get("image"), project.get("image_caption", "项目配图"))
        add_para(doc, f"项目地址：{project.get('url', '')}", size=9, color=(12, 111, 126))
        for paragraph in project.get("paragraphs", []):
            add_para(doc, paragraph)
        add_bullets(doc, project.get("bullets"))
        takeaway = project.get("takeaway")
        if takeaway:
            add_para(doc, "一句话推荐：" + takeaway, size=10.5, bold=True, color=(12, 111, 126), space_after=12)

    observation = article.get("observation", {})
    if observation:
        add_heading(doc, observation.get("title", "本周观察"), 1)
        for paragraph in observation.get("paragraphs", []):
            add_para(doc, paragraph)
        add_bullets(doc, observation.get("bullets"))
        if observation.get("takeaway"):
            add_para(doc, observation["takeaway"], size=10.5, bold=True, color=(12, 111, 126), space_after=12)

    doc.save(out_path)


def sample_article():
    return {
        "title": "本周 GitHub 热门内容推荐",
        "subtitle": "AI Agent 正在从会写代码走向会工作",
        "metadata": "数据参考：GitHub Trending 月榜与项目主页。时间：2026 年 5 月 24 日。",
        "cover_image": "",
        "cover_caption": "本周 GitHub 热点封面图",
        "intro": [
            "这一周的 GitHub 热点，表面上看还是 AI Agent 很火，但仔细看会发现重点已经变了。",
            "大家不再只是关心 AI 能不能写代码，而是开始关心它能不能在真实项目里稳定工作。",
        ],
        "summary": [
            {
                "name": "owner/repo",
                "direction": "AI 编程规范",
                "value": "给 Coding Agent 建立行为边界",
                "reason": "适合给 Agent 加一套工作纪律",
            }
        ],
        "projects": [
            {
                "name": "owner/repo",
                "url": "https://github.com/owner/repo",
                "image": "",
                "image_caption": "项目配图",
                "paragraphs": [
                    "这里介绍项目是什么，以及它为什么值得关注。",
                    "这里解释核心思路、方向、创新点和适合人群。",
                ],
                "bullets": ["要点一", "要点二"],
                "takeaway": "一句自然、准确的推荐理由。",
            }
        ],
        "observation": {
            "title": "本周观察",
            "paragraphs": ["这里总结 5 个项目背后的共同趋势。"],
            "bullets": ["趋势一", "趋势二"],
            "takeaway": "AI 编程正在从生成代码进入组织工作的阶段。",
        },
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("article_json", nargs="?")
    parser.add_argument("--out", default=None, help="Output DOCX path. Defaults to previous-week filename in the configured output folder.")
    parser.add_argument("--write-sample")
    parser.add_argument("--print-default-out", action="store_true")
    args = parser.parse_args()

    if args.print_default_out:
        print(default_output_path())
        return

    if args.write_sample:
        Path(args.write_sample).write_text(json.dumps(sample_article(), ensure_ascii=False, indent=2), encoding="utf-8")
        print(args.write_sample)
        return

    if not args.article_json:
        raise SystemExit("usage: build_github_weekly_docx.py article.json [--out output.docx]")
    article = json.loads(Path(args.article_json).read_text(encoding="utf-8"))
    out_path = Path(args.out) if args.out else default_output_path()
    build_docx(article, out_path)
    print(out_path)


if __name__ == "__main__":
    main()
